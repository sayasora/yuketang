from docx import Document
from docx.shared import Inches
import requests
from PIL import Image
from read_img import get_text
import io
import os

def get_url(a,b,c):
    # 需要获取登陆账户 获取cookie 否则无法爬取
    headers = {
        'authority': 'www.yuketang.cn',
        'pragma': 'no-cache',
        'cache-control': 'no-cache',
        'accept': 'application/json, text/plain, */*',
        'sec-fetch-dest': 'empty',
        'xtbz': 'ykt',
        'university-id': '0',
        'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/80.0.3987.132 Safari/537.36',
        'xt-agent': 'web',
        'sec-fetch-site': 'same-origin',
        'sec-fetch-mode': 'cors',
        'referer': 'https://www.yuketang.cn/v2/web/quizSummary/3545590/878580',
        'accept-language': 'zh-CN,zh;q=0.9',
        修改此处'cookie': '',
    }

    params = (
        ('classroom_id', str(a)),
        ('quiz_id', str(b)),
        ('problem_id', str(c)),
    )

    response = requests.get('https://www.yuketang.cn/v2/api/web/quiz/problem_shape', headers=headers, params=params)
    text = response.text
    return text

def get_aq(text):
    text_all = eval(text.replace('true', 'True').replace('false', 'False'))
    urls = text_all['data']['Shapes']
    dicts = {'0': '', '1': 'A：', '2': 'B：', '3': 'C：', '4': 'D：', '5': '图片', '6': '图片', '7': '图片', '8': '图片', '9': '图片', '10': '图片', '11': '图片'}
    for i in range(len(urls)):
        url = urls[i].get('URL')
        text = get_text(url)
        if text['state'] == 1:
            play_text = text['content']
            print(play_text)
            document.add_paragraph(dicts[str(i)] + play_text)
        else:
            url = text['url']
            response = requests.get(url)
            a = response.content
            img = Image.open(io.BytesIO(a))
            img2 = transparence2white(img)
            img2.save('bar3.png')
            pic = document.add_paragraph(dicts[str(i)])
            pic.add_picture('bar3.png', width=Inches(1))
            os.remove('bar3.png')
    document.add_paragraph("答案： " + text_all['data']['Answer'])

def transparence2white(img):
#     img=img.convert('RGBA')  # 此步骤是将图像转为灰度(RGBA表示4x8位像素，带透明度掩模的真彩色；CMYK为4x8位像素，分色等)，可以省略
    sp=img.size
    width=sp[0]
    height=sp[1]
    print(sp)
    for yh in range(height):
        for xw in range(width):
            dot=(xw,yh)
            color_d=img.getpixel(dot)  # 与cv2不同的是，这里需要用getpixel方法来获取维度数据
            if(color_d[3]==0):
                color_d=(255,255,255,255)
                img.putpixel(dot,color_d)  # 赋值的方法是通过putpixel
    return img



if __name__ == '__main__':
    document = Document()

    classroom_id = 3545590
    # 试题id
    quiz_id = 878580
    # 第一道题id
    problem_id = 19068856
    # 最后一道题id
    problem_id2 = 19068859
    # 题号
    num = 1

    for i in range(problem_id, problem_id2 + 1):
        print()
        print('题号：' + str(num) + '  参数'+str([classroom_id, quiz_id, i]))
        text = get_url(classroom_id, quiz_id, i)
        document.add_paragraph('题号：' + str(num) + '  参数' + str([classroom_id, quiz_id, i]))
        get_aq(text)
        num = num + 1
        document.add_paragraph()

    document.save('测试7.docx')